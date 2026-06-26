"""
ml/parser.py - Resume Parser
==============================
This is the NLP core that extracts structured information from raw resume text.

PIPELINE:
1. Extract raw text from PDF or DOCX file
2. Run spaCy NER (Named Entity Recognition) to find names, organizations
3. Use regex patterns to extract email, phone
4. Use skill matching against a known skills database
5. Parse education and experience sections

KEY NLP CONCEPTS:
- NER (Named Entity Recognition): ML model that labels words as PERSON, ORG, DATE, etc.
- Regex: Pattern matching for structured data like emails and phone numbers
- Tokenization: Splitting text into words/sentences for processing
- Lemmatization: Reducing words to root form (running → run, skills → skill)
"""

import re
import fitz          # PyMuPDF — for PDF extraction
import docx          # python-docx — for DOCX extraction
import spacy
import nltk
from typing import Optional, List, Dict, Any
from pathlib import Path

# Download required NLTK data (runs only when needed)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
    
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger', quiet=True)


# Load spaCy English model
# en_core_web_sm = small English model (fast, good enough for resume parsing)
# Install with: python -m spacy download en_core_web_sm
try:
    nlp = spacy.load("en_core_web_sm")
    print("✅ spaCy model loaded")
except OSError:
    print("⚠️ spaCy model not found. Run: python -m spacy download en_core_web_sm")
    nlp = None


# ============================================================
# Skills Database
# ============================================================
# A comprehensive list of technical and soft skills to match against.
# In production, you could load this from a database or external file.

SKILLS_DATABASE = {
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "ruby",
    "go", "golang", "rust", "swift", "kotlin", "scala", "r", "matlab",
    "php", "perl", "shell", "bash", "powershell", "dart", "lua",
    
    # Web Frameworks
    "react", "react.js", "reactjs", "angular", "vue", "vue.js", "vuejs",
    "node", "node.js", "nodejs", "express", "django", "flask", "fastapi",
    "spring", "spring boot", "laravel", "rails", "ruby on rails",
    "next.js", "nextjs", "nuxt", "gatsby", "svelte",
    
    # Databases
    "mysql", "postgresql", "postgres", "mongodb", "redis", "elasticsearch",
    "cassandra", "dynamodb", "sqlite", "oracle", "sql server", "mariadb",
    "neo4j", "influxdb", "firebase",
    
    # Cloud & DevOps
    "aws", "gcp", "azure", "docker", "kubernetes", "k8s", "terraform",
    "ansible", "jenkins", "gitlab ci", "github actions", "ci/cd",
    "linux", "unix", "nginx", "apache", "microservices",
    
    # ML/AI
    "machine learning", "deep learning", "neural networks", "nlp",
    "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn",
    "pandas", "numpy", "matplotlib", "seaborn", "spark", "hadoop",
    "bert", "gpt", "transformers", "opencv", "yolo",
    
    # Data & Analytics
    "sql", "data analysis", "data science", "statistics", "tableau",
    "power bi", "excel", "etl", "data warehousing", "airflow",
    
    # Tools & Practices
    "git", "github", "gitlab", "bitbucket", "jira", "confluence",
    "agile", "scrum", "kanban", "tdd", "rest api", "graphql", "grpc",
    "microservices", "api design", "system design",
    
    # Soft Skills
    "leadership", "communication", "teamwork", "problem solving",
    "project management", "critical thinking", "collaboration",
    "time management", "mentoring", "documentation",
}


# ============================================================
# File Text Extraction
# ============================================================

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts all text from a PDF file using PyMuPDF.
    
    PyMuPDF (fitz) is faster and more accurate than other PDF libraries.
    It handles multi-column layouts, tables, and embedded fonts better.
    
    Parameters:
        file_path: Path to the PDF file
    
    Returns:
        Extracted text as a single string
    """
    try:
        text = ""
        # fitz.open() loads the PDF into memory
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            # get_text() extracts text from one page
            page_text = page.get_text("text")  # "text" mode = plain text
            text += page_text + "\n"
        
        doc.close()
        return text.strip()
    
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts all text from a DOCX file using python-docx.
    
    DOCX files are actually ZIP archives containing XML.
    python-docx parses the XML and gives us structured access to paragraphs.
    
    Parameters:
        file_path: Path to the DOCX file
    
    Returns:
        Extracted text as a single string
    """
    try:
        doc = docx.Document(file_path)
        
        # Each paragraph is a block of text in the document
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                paragraphs.append(para.text)
        
        # Also extract text from tables (some resumes use table layouts)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return "\n".join(paragraphs)
    
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    Auto-detects file type and extracts text.
    
    Dispatches to the correct extractor based on file extension.
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif filename_lower.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {filename}. Supported: PDF, DOCX, TXT")


# ============================================================
# Information Extraction Functions
# ============================================================

def extract_email(text: str) -> Optional[str]:
    """
    Extracts email address using regex.
    
    Regex pattern breakdown:
    [a-zA-Z0-9._%+-]+  = one or more valid email username characters
    @                  = literal @ symbol
    [a-zA-Z0-9.-]+     = domain name
    \.                 = literal dot
    [a-zA-Z]{2,}       = top-level domain (com, org, net, etc.)
    """
    email_pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
    matches = re.findall(email_pattern, text)
    return matches[0] if matches else None


def extract_phone(text: str) -> Optional[str]:
    """
    Extracts phone number using regex.
    Handles various formats: +91-9876543210, (555) 123-4567, 9876543210
    """
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # International
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',                    # US format
        r'\d{10}',                                                    # 10 digits
        r'\+\d{10,13}',                                              # +91XXXXXXXXXX
    ]
    
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Clean up the phone number
            phone = re.sub(r'[^\d+\-\(\)\s]', '', matches[0])
            return phone.strip()
    
    return None


def extract_name(text: str) -> Optional[str]:
    """
    Extracts candidate name using spaCy NER.
    
    spaCy's PERSON entity type identifies human names.
    We take the first PERSON entity found (usually at the top of the resume).
    
    Fallback: Take the first non-empty line (usually the name on a resume).
    """
    if nlp:
        # Process first 500 characters (name is almost always at the top)
        doc = nlp(text[:500])
        
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                name = ent.text.strip()
                # Filter out very short or very long names
                if 2 <= len(name.split()) <= 4 and len(name) > 3:
                    return name
    
    # Fallback: First non-empty line
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        first_line = lines[0]
        # Likely a name if it's 1-4 words and no special characters
        if re.match(r'^[A-Za-z\s\.]+$', first_line) and len(first_line.split()) <= 4:
            return first_line
    
    return None


def extract_skills(text: str) -> List[str]:
    """
    Extracts skills from resume text by matching against our skills database.
    
    APPROACH:
    1. Convert text to lowercase for case-insensitive matching
    2. Check for each skill in our database as a whole word/phrase
    3. Return matched skills (deduplicated)
    
    WHY WHOLE WORD MATCHING?
    - Without word boundaries, "c" would match in "because", "science", etc.
    - re.escape() safely handles special regex characters in skill names
    """
    text_lower = text.lower()
    found_skills = set()
    
    for skill in SKILLS_DATABASE:
        # Use word boundary (\b) to match whole words only
        # re.escape handles skills with special chars like "c++" or "node.js"
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            # Store in title case for nice display
            found_skills.add(skill.title())
    
    return sorted(list(found_skills))


def extract_education(text: str) -> List[Dict]:
    """
    Extracts education entries using pattern matching.
    
    Looks for:
    - Degree keywords (B.Tech, MBA, Bachelor, Master, PhD, etc.)
    - University/College names (via spaCy ORG entities)
    - Years (4-digit numbers in range 1980-2030)
    """
    education = []
    
    # Common degree patterns
    degree_patterns = [
        r'\b(B\.?Tech|B\.?E|B\.?Sc|B\.?S|Bachelor[\'s]?\s+of\s+\w+)\b',
        r'\b(M\.?Tech|M\.?E|M\.?Sc|M\.?S|Master[\'s]?\s+of\s+\w+|MBA)\b',
        r'\b(Ph\.?D|Doctor\s+of\s+\w+)\b',
        r'\b(Diploma|Certificate|Associate[\'s]?\s+Degree)\b',
        r'\b(High\s+School|12th|10th|Secondary|Senior\s+Secondary)\b',
    ]
    
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        for pattern in degree_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                entry = {
                    "degree": match.group(0),
                    "institution": None,
                    "year": None,
                    "gpa": None
                }
                
                # Look for year in this line and next 2 lines
                context = " ".join(lines[max(0, i-1):min(len(lines), i+3)])
                year_match = re.search(r'\b(19|20)\d{2}\b', context)
                if year_match:
                    entry["year"] = year_match.group(0)
                
                # Look for institution in context
                if nlp:
                    doc = nlp(context)
                    for ent in doc.ents:
                        if ent.label_ == "ORG" and len(ent.text) > 3:
                            entry["institution"] = ent.text
                            break
                
                # Look for GPA
                gpa_match = re.search(r'GPA[:\s]+(\d+\.?\d*)', context, re.IGNORECASE)
                if gpa_match:
                    entry["gpa"] = gpa_match.group(1)
                
                education.append(entry)
                break
    
    return education[:5]  # Return max 5 education entries


def extract_experience(text: str) -> List[Dict]:
    """
    Extracts work experience entries.
    
    Looks for sections containing experience keywords and extracts:
    - Job titles (Engineer, Manager, Developer, etc.)
    - Company names (via spaCy ORG entities)  
    - Duration (date ranges)
    """
    experience = []
    
    # Split text into sections
    sections = re.split(
        r'\n(?=(?:EXPERIENCE|WORK\s+HISTORY|EMPLOYMENT|PROFESSIONAL\s+EXPERIENCE))',
        text, flags=re.IGNORECASE
    )
    
    # Job title patterns
    job_title_patterns = [
        r'\b(Senior\s+)?(?:Software|Full[\s-]Stack|Frontend|Backend|Data|ML|AI|'
        r'DevOps|Cloud|Mobile|iOS|Android|QA|Security|Product|Project|'
        r'Solutions)\s+(?:Engineer|Developer|Architect|Manager|Analyst|Scientist|Lead)\b',
        r'\b(?:CTO|CEO|VP|Director|Head|Lead|Principal|Staff)\s+of?\s+\w+\b',
        r'\b(?:Intern|Junior|Mid[\s-]level|Senior|Principal)\s+\w+\b',
    ]
    
    # Date range pattern: "Jan 2020 - Dec 2022" or "2020 - 2022"
    date_pattern = (
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|'
        r'March|April|June|July|August|September|October|November|December)?'
        r'\s*\d{4}\s*[-–—]\s*'
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|'
        r'March|April|June|July|August|September|October|November|December)?'
        r'\s*(?:\d{4}|Present|Current|Now)'
    )
    
    lines = text.split('\n')
    
    for i, line in enumerate(lines):
        for pattern in job_title_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                entry = {
                    "title": match.group(0).strip(),
                    "company": None,
                    "duration": None,
                    "description": None
                }
                
                # Get context around this line
                context = " ".join(lines[max(0, i-1):min(len(lines), i+4)])
                
                # Extract company using spaCy
                if nlp:
                    doc = nlp(context[:300])
                    for ent in doc.ents:
                        if ent.label_ == "ORG" and len(ent.text) > 2:
                            entry["company"] = ent.text
                            break
                
                # Extract duration
                duration_match = re.search(date_pattern, context, re.IGNORECASE)
                if duration_match:
                    entry["duration"] = duration_match.group(0).strip()
                
                # Get description (next few lines)
                desc_lines = [l.strip() for l in lines[i+1:i+4] if l.strip()]
                if desc_lines:
                    entry["description"] = " ".join(desc_lines[:2])
                
                experience.append(entry)
                break
    
    return experience[:10]  # Return max 10 experience entries


def calculate_experience_years(experience: List[Dict]) -> float:
    """
    Estimates total years of experience from parsed experience entries.
    """
    total_years = 0.0
    current_year = 2024
    
    for exp in experience:
        duration = exp.get("duration", "")
        if not duration:
            continue
        
        # Find all years mentioned
        years = re.findall(r'\b(19|20)(\d{2})\b', duration)
        years = [int(f"{y[0]}{y[1]}") for y in years]
        
        # Check for "Present/Current"
        has_present = bool(re.search(r'Present|Current|Now', duration, re.IGNORECASE))
        
        if len(years) >= 1:
            start_year = min(years)
            end_year = current_year if has_present else (max(years) if len(years) >= 2 else start_year + 1)
            total_years += max(0, end_year - start_year)
    
    return min(total_years, 40.0)  # Cap at 40 years


# ============================================================
# Main Parser Function
# ============================================================

def parse_resume(file_path: str, filename: str) -> Dict[str, Any]:
    """
    Main function that orchestrates the full resume parsing pipeline.
    
    Steps:
    1. Extract raw text from file
    2. Run all extraction functions
    3. Return structured data
    
    Parameters:
        file_path: Path to the saved resume file
        filename: Original filename (used to detect file type)
    
    Returns:
        Dictionary with all extracted fields
    """
    # Step 1: Extract raw text
    raw_text = extract_text_from_file(file_path, filename)
    
    if not raw_text or len(raw_text) < 50:
        raise ValueError("Could not extract meaningful text from the resume file")
    
    # Step 2: Extract structured information
    experience_list = extract_experience(raw_text)
    
    parsed = {
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": extract_skills(raw_text),
        "education": extract_education(raw_text),
        "experience": experience_list,
        "total_experience_years": calculate_experience_years(experience_list),
        "raw_text": raw_text
    }
    
    return parsed
